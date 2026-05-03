from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level, color_hex="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.size = Pt(11)
    return p


def add_table_header(table, headers, bg="1F3864"):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)


def add_table_row(table, values, bg=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        if bg:
            set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    return row


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# DOCUMENTO 1: INFORME FINAL
# ─────────────────────────────────────────────

def generar_informe():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── Portada ──────────────────────────────
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("INFORME FINAL DE PASANTÍA")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("1F3864")

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run("Módulo de Sistemas – Aplicación HRCATCH 2026")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor.from_string("2E74B5")

    doc.add_paragraph()
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.add_run("Hospital Universitario San Rafael").bold = True

    doc.add_paragraph()
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_p.add_run("Año 2026")

    page_break(doc)

    # ── 1. Introducción ──────────────────────
    add_heading(doc, "1. Introducción", 1)
    add_body(doc,
        "El presente informe describe las actividades realizadas durante la pasantía en el "
        "Hospital Universitario San Rafael (HUSRT), enfocándose exclusivamente en el desarrollo "
        "del Módulo de Sistemas de la aplicación HRCATCH 2026. Este módulo tiene como propósito "
        "centralizar y digitalizar la gestión del inventario tecnológico de la institución, "
        "permitiendo el control integral del ciclo de vida de los equipos de cómputo e "
        "infraestructura tecnológica.")

    add_body(doc,
        "La aplicación HRCATCH 2026 es una plataforma web desarrollada con Angular en el "
        "frontend y una API REST en el backend, apoyada en la librería de componentes PrimeNG "
        "para la interfaz de usuario. El sistema está orientado a múltiples roles: "
        "Superadministrador, Administrador de Sistemas y Usuario de Sistemas, cada uno con "
        "niveles de acceso diferenciados.")

    add_body(doc,
        "En el Módulo de Sistemas se trabajaron dos grandes áreas: el Card de Inventario, "
        "que concentra toda la funcionalidad de gestión de equipos, y el Módulo de "
        "Parametrización SIS, que permite configurar los catálogos y parámetros base del "
        "sistema. Ambas áreas son el objeto de análisis de este informe.")

    page_break(doc)

    # ── 2. Informe de Desarrollo de Actividades ──
    add_heading(doc, "2. Informe de Desarrollo de Actividades", 1)

    add_body(doc,
        "El desarrollo del módulo se realizó de forma iterativa, integrando las vistas de "
        "frontend con los endpoints del backend a medida que cada funcionalidad fue acordada "
        "con el equipo. A continuación se describe en detalle cada componente implementado.")

    # 2.1 Card de Inventario
    add_heading(doc, "2.1 Card de Inventario", 2)
    add_body(doc,
        "El Card de Inventario es el punto de entrada principal a la gestión de equipos "
        "tecnológicos. Desde el dashboard del módulo de Sistemas, el card presenta cinco "
        "accesos directos: Tipos de Equipo, Equipos en Bodega, Dados de Baja, Por Servicio "
        "y Por Sede. Cada acceso redirige a una vista especializada que comparte la misma "
        "lógica de negocio pero ofrece perspectivas de navegación distintas.")

    # 2.1.1
    add_heading(doc, "2.1.1 Vista de Tipos de Equipo", 3)
    add_body(doc,
        "Esta vista presenta una cuadrícula de tarjetas (cards), donde cada tarjeta representa "
        "un tipo de equipo registrado en el sistema (por ejemplo, Computador de Escritorio, "
        "Portátil, Impresora, Switch, etc.). Cada tarjeta muestra el nombre del tipo y el "
        "conteo de equipos registrados bajo esa categoría.")
    add_body(doc, "Funcionalidades disponibles:")
    add_bullet(doc, "Búsqueda en tiempo real por nombre de tipo de equipo.")
    add_bullet(doc, "Acceso directo a la lista de equipos filtrada por tipo al hacer clic en una tarjeta.")
    add_bullet(doc, "Botón de creación rápida de nuevo equipo.")
    add_bullet(doc, "Exportación del inventario a Excel con filtros: todos los equipos, solo activos, en bodega o inactivos.")

    # 2.1.2
    add_heading(doc, "2.1.2 Lista de Equipos (por Tipo, Sede o Servicio)", 3)
    add_body(doc,
        "Una vez el usuario selecciona un tipo de equipo, una sede o un servicio, accede a "
        "una tabla paginada (15 ítems por página) que lista todos los equipos correspondientes. "
        "La tabla incluye: ID, nombre, tipo, marca, modelo, serial, ubicación y estado (Activo/Inactivo).")
    add_body(doc, "Acciones disponibles por equipo desde el menú contextual:")
    add_bullet(doc, "Ver Detalle: abre un modal con toda la información del equipo.")
    add_bullet(doc, "Editar: abre el formulario de edición pre-cargado con los datos actuales.")
    add_bullet(doc, "Plan de Mantenimiento: programa el intervalo de mantenimiento (mensual a anual) con fecha de inicio.")
    add_bullet(doc, "Reporte de Entrega: genera el formulario de reporte de entrega del equipo.")
    add_bullet(doc, "Ver Reportes: lista todos los reportes generados para el equipo.")
    add_bullet(doc, "Ver Historial: muestra el historial de cambios del equipo.")
    add_bullet(doc, "Enviar a Bodega: pasa el equipo al estado de almacenamiento, requiriendo motivo.")
    add_bullet(doc, "Dar de Baja: retira el equipo del inventario activo, requiriendo justificación y contraseña.")
    add_bullet(doc, "Ver Hoja de Vida (HV): navega a la ficha técnica completa del equipo.")
    add_body(doc,
        "La tabla cuenta con búsqueda global que filtra por nombre, marca, modelo, serial, "
        "código de inventario, ubicación e ID, así como un filtro por estado (Activo/Inactivo).")

    # 2.1.3
    add_heading(doc, "2.1.3 Creación y Edición de Equipos", 3)
    add_body(doc,
        "El formulario de creación/edición de equipos es un modal organizado en secciones "
        "que guía al usuario a través de todos los datos necesarios:")
    add_bullet(doc, "Información Básica: nombre, marca, modelo, serial, placa de inventario, código, año de adquisición.")
    add_bullet(doc, "Ubicación: ubicación general y ubicación específica.")
    add_bullet(doc, "Mantenimiento: indica si requiere mantenimiento, frecuencia (mensual, trimestral, semestral, anual) y fechas calculadas automáticamente.")
    add_bullet(doc, "Configuración de Red: si el equipo es gestionable, número de puertos, direccionamiento VLAN.")
    add_bullet(doc, "Asignaciones: sede (que filtra dinámicamente los servicios disponibles), servicio, tipo de equipo (que define qué campos de hoja de vida se mostrarán) y usuario responsable.")
    add_bullet(doc, "Hoja de Vida (solo en creación): campos técnicos dinámicos según el tipo de equipo configurado en parametrización (procesador, RAM, disco, OS, suite ofimática, IP, MAC, proveedor, fechas de compra e instalación, costo, contrato, tipo de adquisición, tipo de uso, observaciones).")
    add_body(doc,
        "El formulario aplica validación en tiempo real, deshabilitando el botón de guardado "
        "mientras haya campos obligatorios incompletos, y muestra un spinner de carga durante "
        "el proceso de guardado.")

    # 2.1.4
    add_heading(doc, "2.1.4 Hoja de Vida del Equipo", 3)
    add_body(doc,
        "La Hoja de Vida (HV) es la ficha técnica completa de cada equipo. Se accede desde "
        "el menú de acciones de cualquier equipo. Está organizada en las siguientes secciones:")
    add_bullet(doc, "Datos básicos del equipo (lectura desde el registro principal).")
    add_bullet(doc, "Especificaciones técnicas editables: procesador, RAM, disco duro, tóner/cartucho, sistema operativo, suite ofimática, IP, MAC. Los campos visibles dependen de la configuración del tipo de equipo.")
    add_bullet(doc, "Información de adquisición: proveedor, fecha de compra, fecha de instalación, costo, número de contrato, tipo de adquisición (compra directa, convenio, donación, comodato).")
    add_bullet(doc, "Tipo de uso y observaciones generales.")
    add_bullet(doc, "Foto del equipo: permite subir, visualizar y reemplazar una imagen del equipo.")
    add_bullet(doc, "Documentos adjuntos: permite agregar documentos clasificados por tipo, con subida de archivos, descarga y eliminación.")
    add_bullet(doc, "Historial de ediciones: registro completo de todos los cambios realizados a la HV, con usuario responsable, fecha y tipo de cambio.")
    add_body(doc,
        "La edición de la HV está restringida a usuarios con rol Administrador de Sistemas "
        "o Superadministrador. Los usuarios de solo lectura pueden visualizar toda la información.")

    # 2.1.5
    add_heading(doc, "2.1.5 Gestión del Ciclo de Vida: Bodega y Baja", 3)
    add_body(doc,
        "El sistema contempla un ciclo de vida completo para cada equipo con tres estados "
        "principales: Activo, En Bodega y Dado de Baja.")
    add_body(doc, "Enviar a Bodega:")
    add_bullet(doc, "Disponible desde el menú de acciones de cualquier equipo activo.")
    add_bullet(doc, "Requiere ingresar un motivo obligatorio para el registro del cambio.")
    add_bullet(doc, "El equipo desaparece de la vista activa y aparece en la vista 'Equipos en Bodega'.")
    add_body(doc, "Dar de Baja:")
    add_bullet(doc, "Proceso formal con diálogo de confirmación que solicita: justificación detallada, inventario de accesorios reutilizables, identificación del usuario y contraseña de seguridad.")
    add_bullet(doc, "El equipo pasa al estado 'Dado de Baja' y se puede descargar el acta en PDF.")
    add_body(doc, "Reactivar desde Bodega:")
    add_bullet(doc, "Desde la vista de Equipos en Bodega, el superadministrador puede reactivar un equipo.")
    add_bullet(doc, "El proceso permite opcionalmente reasignar la ubicación del equipo al reactivarlo.")

    # 2.1.6
    add_heading(doc, "2.1.6 Vistas por Sede y por Servicio", 3)
    add_body(doc,
        "Adicionalmente a la navegación por tipo de equipo, el sistema ofrece dos perspectivas "
        "alternativas de consulta del inventario:")
    add_body(doc, "Por Sede:")
    add_bullet(doc, "Muestra todas las sedes (instalaciones físicas) del hospital en tarjetas.")
    add_bullet(doc, "Cada tarjeta indica el nombre de la sede, el número de equipos asignados y la cantidad de servicios.")
    add_bullet(doc, "Filtros: todas las sedes, sedes con equipos asignados, sedes sin equipos.")
    add_bullet(doc, "Al seleccionar una sede se accede a la lista de equipos en esa ubicación con todas las funcionalidades del menú de acciones.")
    add_body(doc, "Por Servicio:")
    add_bullet(doc, "Similar a la vista por sede, pero organiza el inventario por departamento o servicio hospitalario.")
    add_bullet(doc, "Filtros: todos los servicios, servicios con equipos, servicios sin equipos.")
    add_bullet(doc, "Facilita la identificación de equipos por área funcional (por ejemplo, Urgencias, Laboratorio, Radiología).")

    page_break(doc)

    # 2.2 Módulo de Parametrización
    add_heading(doc, "2.2 Módulo de Parametrización SIS", 2)
    add_body(doc,
        "El Módulo de Parametrización SIS es el área de configuración del sistema, visible "
        "únicamente para usuarios con rol Administrador de Sistemas o Superadministrador. "
        "Permite gestionar los catálogos base que alimentan todo el módulo de inventario. "
        "Se accede desde el dashboard principal de Sistemas a través del card 'Parametrización'.")

    # 2.2.1
    add_heading(doc, "2.2.1 Administración de Tipos de Equipo", 3)
    add_body(doc,
        "Esta sección es el núcleo de la parametrización. Permite definir y gestionar los "
        "tipos de equipo que clasifican todo el inventario. Se implementó con una tabla "
        "PrimeNG DataTable con ordenamiento, paginación configurable (10, 25 ó 50 filas) "
        "y búsqueda global.")
    add_body(doc, "Datos gestionados por tipo de equipo:")
    add_bullet(doc, "Nombre del tipo.")
    add_bullet(doc, "Materiales consumibles requeridos para el mantenimiento.")
    add_bullet(doc, "Herramientas necesarias para el mantenimiento.")
    add_bullet(doc, "Tiempo estimado de mantenimiento (en minutos).")
    add_bullet(doc, "Repuestos mínimos recomendados.")
    add_bullet(doc, "Tipo de actividad de mantenimiento.")
    add_bullet(doc, "Estado (Activo/Inactivo). No se puede desactivar un tipo si tiene equipos activos asignados.")
    add_body(doc,
        "Una de las funcionalidades más relevantes de esta sección es la configuración de "
        "la Hoja de Vida (HV). Por cada tipo de equipo se pueden habilitar o deshabilitar "
        "individualmente los campos que aparecerán en la hoja de vida de los equipos de ese "
        "tipo. Los campos están organizados en cuatro grupos:")

    # tabla de campos HV
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    add_table_header(t, ["Grupo", "Campos configurables"])
    grupos = [
        ("Red", "Dirección IP, Dirección MAC"),
        ("Hardware", "Procesador, Memoria RAM, Disco Duro, Tóner/Cartucho"),
        ("Software", "Sistema Operativo, Suite Ofimática"),
        ("Gestión", "Nombre de usuario, Tipo de uso, Fecha de adquisición, Observaciones"),
    ]
    for i, (g, c) in enumerate(grupos):
        bg = "EBF3FB" if i % 2 == 0 else None
        add_table_row(t, [g, c], bg=bg)

    doc.add_paragraph()
    add_body(doc,
        "Esta configuración dinámica permite que el formulario de creación de equipos y la "
        "hoja de vida solo muestren los campos relevantes para cada tipo, evitando información "
        "innecesaria y simplificando la experiencia de usuario.")

    # 2.2.2
    add_heading(doc, "2.2.2 Gestión de Protocolos de Mantenimiento Preventivo", 3)
    add_body(doc,
        "Cada tipo de equipo puede tener asociados protocolos de mantenimiento preventivo, "
        "que son los pasos o procedimientos que el técnico debe seguir al realizar una "
        "intervención preventiva. Esta sección permite:")
    add_bullet(doc, "Seleccionar un tipo de equipo desde un dropdown.")
    add_bullet(doc, "Ver la lista paginada de todos los protocolos (pasos) definidos para ese tipo.")
    add_bullet(doc, "Agregar nuevos pasos al protocolo mediante un formulario de texto.")
    add_bullet(doc, "Editar el texto de pasos existentes.")
    add_bullet(doc, "Activar o desactivar pasos individuales del protocolo.")
    add_body(doc,
        "Los protocolos activos son los que se presentarán al técnico al momento de ejecutar "
        "un mantenimiento preventivo sobre un equipo del tipo correspondiente, garantizando "
        "estandarización en los procedimientos de mantenimiento.")

    page_break(doc)

    # ── 3. Objetivos ──────────────────────────
    add_heading(doc, "3. Objetivos", 1)

    add_heading(doc, "3.1 Objetivo General", 2)
    add_body(doc,
        "Desarrollar el Módulo de Sistemas de la aplicación HRCATCH 2026 del Hospital "
        "Universitario San Rafael, implementando una solución integral para la gestión "
        "del inventario de equipos tecnológicos que permita el registro, seguimiento y "
        "control del ciclo de vida de los activos de cómputo e infraestructura.")

    add_heading(doc, "3.2 Objetivos Específicos", 2)
    add_bullet(doc, "Implementar el módulo de inventario de equipos con vistas de navegación por tipo, sede y servicio.")
    add_bullet(doc, "Desarrollar el formulario de creación y edición de equipos con validación dinámica según el tipo de equipo seleccionado.")
    add_bullet(doc, "Implementar la gestión del ciclo de vida de los equipos: estado activo, envío a bodega, reactivación y proceso de baja.")
    add_bullet(doc, "Desarrollar la hoja de vida del equipo con soporte para adjuntar fotografías y documentos, y con trazabilidad de cambios.")
    add_bullet(doc, "Construir el módulo de parametrización que permita configurar tipos de equipo, los campos de hoja de vida aplicables y los protocolos de mantenimiento preventivo.")
    add_bullet(doc, "Integrar el sistema de roles y permisos de la aplicación para controlar el acceso a las funcionalidades según el perfil del usuario.")

    page_break(doc)

    # ── 4. Logros Alcanzados ──────────────────
    add_heading(doc, "4. Logros Alcanzados", 1)
    add_body(doc,
        "Al finalizar el período de pasantía, se lograron implementar y poner en funcionamiento "
        "los siguientes entregables:")

    logros = [
        ("Inventario de equipos funcional",
         "Se desarrolló la gestión completa de equipos tecnológicos con CRUD operativo, "
         "paginación, búsqueda y filtros por estado. Las vistas por tipo, sede y servicio "
         "permiten múltiples perspectivas de consulta sobre el mismo inventario."),
        ("Ciclo de vida completo del equipo",
         "Se implementaron los flujos de envío a bodega (con motivo), baja de equipos "
         "(con justificación formal y contraseña de seguridad), descarga del acta de baja "
         "en PDF y reactivación de equipos desde bodega."),
        ("Hoja de vida dinámica",
         "Se desarrolló la ficha técnica de cada equipo con campos que se muestran u "
         "ocultan según la configuración del tipo de equipo, soporte para subida de fotos "
         "y documentos adjuntos, y un historial de cambios trazable por usuario y fecha."),
        ("Módulo de parametrización",
         "Se implementó la administración de tipos de equipo con control granular de qué "
         "campos de hoja de vida aplican a cada tipo, y la gestión de protocolos de "
         "mantenimiento preventivo por tipo, con activación/desactivación individual de pasos."),
        ("Exportación de inventario",
         "Se integró la funcionalidad de exportar el inventario en formato Excel, con "
         "filtros por estado del equipo (todos, activos, en bodega, inactivos)."),
        ("Plan de mantenimiento",
         "Se implementó el diálogo de programación de mantenimiento que permite definir "
         "la frecuencia (mensual a anual) y la fecha de inicio, generando automáticamente "
         "las fechas de mantenimiento previstas."),
        ("Integración de roles",
         "Se implementó el control de acceso basado en roles para todas las funcionalidades: "
         "el Superadministrador y el Administrador de Sistemas tienen acceso completo, "
         "mientras que el Usuario de Sistemas tiene acceso de solo lectura y consulta. "
         "El card de Parametrización es exclusivo para administradores."),
    ]

    for titulo_l, desc in logros:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{titulo_l}: ")
        run_t.bold = True
        run_t.font.size = Pt(11)
        run_d = p.add_run(desc)
        run_d.font.size = Pt(11)

    page_break(doc)

    # ── 5. Lecciones Aprendidas ──────────────
    add_heading(doc, "5. Lecciones Aprendidas", 1)

    add_heading(doc, "5.1 Técnicas", 2)
    add_bullet(doc, "Profundización en Angular con manejo de formularios reactivos, validación dinámica y comunicación entre componentes padre-hijo mediante @Input/@Output y servicios compartidos.")
    add_bullet(doc, "Uso avanzado de PrimeNG: DataTable con paginación, ordenamiento y filtros globales; diálogos modales; menús contextuales; y notificaciones Toast.")
    add_bullet(doc, "Implementación de lógica de roles y permisos en el frontend usando guardas de rutas y directivas condicionales (@if) basadas en el perfil del usuario autenticado.")
    add_bullet(doc, "Consumo de APIs REST con manejo de estados de carga, errores y respuestas asíncronas mediante RxJS (Observables, switchMap, forkJoin).")
    add_bullet(doc, "Diseño de formularios con campos dinámicos cuya visibilidad depende de datos configurados en el backend, logrando una UI adaptativa sin hardcodear reglas de negocio.")
    add_bullet(doc, "Exportación de datos en Excel y generación de PDFs como parte de flujos de negocio formales (actas de baja, reportes de entrega).")

    add_heading(doc, "5.2 Profesionales y de Proceso", 2)
    add_bullet(doc, "Aprendizaje sobre la importancia de la trazabilidad en sistemas hospitalarios: toda acción significativa (baja de equipo, edición de HV) requiere registro de usuario, fecha y motivo.")
    add_bullet(doc, "Comprensión del valor de la parametrización como estrategia de diseño: al hacer configurable la visibilidad de campos, el sistema se adapta a diferentes tipos de activos sin necesidad de modificar código.")
    add_bullet(doc, "Trabajo colaborativo en un equipo de desarrollo con múltiples módulos: coordinación de ramas Git, revisión de merges y alineación con las convenciones del proyecto.")
    add_bullet(doc, "Aplicación de criterios de UX orientados a usuarios no técnicos: simplificar flujos complejos (como la baja de equipos) mediante wizards por pasos con validaciones claras.")

    page_break(doc)

    # ── 6. Conclusiones ──────────────────────
    add_heading(doc, "6. Conclusiones", 1)
    add_body(doc,
        "El desarrollo del Módulo de Sistemas de la aplicación HRCATCH 2026 representó un "
        "aporte concreto a la digitalización de los procesos de gestión del inventario "
        "tecnológico del Hospital Universitario San Rafael. Contar con una herramienta "
        "centralizada para registrar, clasificar y hacer seguimiento a los equipos de "
        "cómputo e infraestructura permite a la institución reducir la dependencia de "
        "registros manuales en hojas de cálculo, mejorar la trazabilidad de los activos "
        "y facilitar la toma de decisiones sobre renovación y mantenimiento del parque tecnológico.")

    add_body(doc,
        "La implementación del ciclo de vida completo del equipo (activo → bodega → baja) "
        "otorga al sistema un rigor formal que se alinea con los requerimientos de control "
        "de activos en entidades de salud, donde la trazabilidad y la justificación documental "
        "de cada movimiento son indispensables.")

    add_body(doc,
        "El Módulo de Parametrización SIS demuestra que una buena arquitectura de "
        "configuración dinámica reduce significativamente el costo de mantenimiento del "
        "sistema: al permitir que los administradores definan qué campos son relevantes "
        "para cada tipo de equipo y qué protocolos deben seguirse en cada mantenimiento, "
        "el sistema se mantiene flexible sin requerir cambios en el código fuente.")

    add_body(doc,
        "Finalmente, la pasantía permitió consolidar competencias técnicas en el stack "
        "Angular/PrimeNG y en el diseño de aplicaciones orientadas a procesos institucionales, "
        "reafirmando que el desarrollo de software en entornos de salud exige especial "
        "atención a la usabilidad, la seguridad de acceso y la integridad de los datos.")

    page_break(doc)

    # ── 7. Anexos ─────────────────────────────
    add_heading(doc, "7. Anexos", 1)
    add_body(doc,
        "A continuación se reserva el espacio para incluir las capturas de pantalla y "
        "diagramas que ilustran las funcionalidades descritas en este informe.")

    anexos = [
        "Anexo A – Dashboard principal del Módulo de Sistemas (Cards de navegación)",
        "Anexo B – Vista de Tipos de Equipo (grid de tarjetas con conteos)",
        "Anexo C – Lista de equipos con menú de acciones desplegado",
        "Anexo D – Formulario de creación de equipo (todas las secciones)",
        "Anexo E – Hoja de Vida de un equipo (con foto y documentos adjuntos)",
        "Anexo F – Diálogo de Baja de equipo (formulario de confirmación)",
        "Anexo G – Vista de Equipos en Bodega",
        "Anexo H – Vista Por Sede (grid de sedes con conteos)",
        "Anexo I – Vista Por Servicio (grid de servicios con conteos)",
        "Anexo J – Módulo de Parametrización: tabla de Tipos de Equipo",
        "Anexo K – Configuración de campos de Hoja de Vida por tipo de equipo",
        "Anexo L – Gestión de Protocolos de Mantenimiento Preventivo",
    ]
    for a in anexos:
        p = doc.add_paragraph(a, style="List Bullet")
        p.style.font.size = Pt(11)
        doc.add_paragraph()  # espacio para pegar la captura

    # Guardar
    path = os.path.join(OUTPUT_DIR, "Informe_Final_Sistemas_HRCATCH2026.docx")
    doc.save(path)
    print(f"Informe guardado: {path}")
    return path


# ─────────────────────────────────────────────
# DOCUMENTO 2: MANUAL DE USUARIO (SUPERADMIN)
# ─────────────────────────────────────────────

def generar_manual():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Portada
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MANUAL DE USUARIO – SUPERADMINISTRADOR")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("1F3864")

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run("Módulo de Sistemas – HRCATCH 2026")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor.from_string("2E74B5")

    doc.add_paragraph()
    ins = doc.add_paragraph()
    ins.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ins.add_run("Hospital Universitario San Rafael").bold = True
    doc.add_paragraph()
    fa = doc.add_paragraph()
    fa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fa.add_run("Año 2026")

    page_break(doc)

    # ── Introducción del manual ────────────────
    add_heading(doc, "1. Introducción", 1)
    add_body(doc,
        "Este manual está dirigido al usuario con perfil Superadministrador de la aplicación "
        "HRCATCH 2026. Describe paso a paso cómo utilizar todas las funcionalidades del "
        "Módulo de Sistemas, incluyendo la gestión del inventario de equipos tecnológicos "
        "y la configuración del sistema a través del Módulo de Parametrización SIS.")
    add_body(doc,
        "El Superadministrador tiene acceso completo a todas las funciones: puede crear, "
        "editar, desactivar, enviar a bodega y dar de baja equipos, así como configurar "
        "los catálogos base del sistema (tipos de equipo, campos de hoja de vida y "
        "protocolos de mantenimiento).")

    page_break(doc)

    # ── 2. Acceso al módulo ────────────────────
    add_heading(doc, "2. Acceso al Módulo de Sistemas", 1)
    add_body(doc,
        "Luego de iniciar sesión en la aplicación HRCATCH 2026 con sus credenciales de "
        "Superadministrador, navegue al módulo de Sistemas desde el menú principal. "
        "Será recibido por el dashboard del módulo, que presenta los siguientes cards:")

    cards_info = [
        ("Inventario", "Acceso a la gestión de equipos por tipo, sede y servicio, así como a equipos en bodega y dados de baja."),
        ("Mantenimiento", "Programación y ejecución de mantenimientos de equipos."),
        ("Backups", "Gestión de copias de seguridad (en desarrollo)."),
        ("Indicadores", "Visualización de indicadores del parque tecnológico (en desarrollo)."),
        ("Inventario de Repuestos", "Gestión de piezas y componentes de repuesto."),
        ("Parametrización", "Configuración de tipos de equipo, protocolos y catálogos base. Visible solo para Administrador y Superadministrador."),
    ]
    tb = doc.add_table(rows=1, cols=2)
    tb.style = "Table Grid"
    add_table_header(tb, ["Card", "Descripción"])
    for i, (c, d) in enumerate(cards_info):
        bg = "EBF3FB" if i % 2 == 0 else None
        add_table_row(tb, [c, d], bg=bg)

    doc.add_paragraph()
    page_break(doc)

    # ── 3. Card de Inventario ──────────────────
    add_heading(doc, "3. Card de Inventario", 1)
    add_body(doc,
        "El Card de Inventario agrupa los cinco accesos principales a la gestión de equipos. "
        "Haga clic en cualquiera de los botones de acción dentro del card para navegar a la "
        "sección correspondiente.")

    # 3.1 Tipos de Equipo
    add_heading(doc, "3.1 Tipos de Equipo", 2)
    add_body(doc,
        "Al acceder a esta sección verá una cuadrícula de tarjetas, una por cada tipo de "
        "equipo configurado en el sistema (Computador, Portátil, Impresora, Switch, etc.).")
    add_body(doc, "¿Qué puede hacer aquí?")
    add_bullet(doc, "Buscar un tipo de equipo: use la barra de búsqueda en la parte superior para filtrar las tarjetas por nombre.")
    add_bullet(doc, "Ver cuántos equipos tiene cada tipo: el número aparece en cada tarjeta.")
    add_bullet(doc, "Ir a la lista de equipos de un tipo: haga clic sobre la tarjeta del tipo deseado.")
    add_bullet(doc, "Crear un nuevo equipo: haga clic en el botón 'Nuevo Equipo' (ícono +).")
    add_bullet(doc, "Exportar el inventario a Excel: haga clic en el botón de exportar y seleccione el filtro deseado: Todos, Activos, En Bodega o Inactivos. Se descargará un archivo .xlsx.")

    # 3.2 Lista de Equipos
    add_heading(doc, "3.2 Lista de Equipos", 2)
    add_body(doc,
        "Después de seleccionar un tipo (o al ingresar por Sede o Servicio), verá una tabla "
        "con todos los equipos. La tabla muestra: ID, Nombre, Tipo, Marca, Modelo, Serial, "
        "Ubicación y Estado.")
    add_body(doc, "Controles de la tabla:")
    add_bullet(doc, "Barra de búsqueda: filtra simultáneamente por nombre, marca, modelo, serial, código y ubicación.")
    add_bullet(doc, "Filtro de estado: seleccione Activo, Inactivo o Todos desde el desplegable.")
    add_bullet(doc, "Paginación: navegue entre páginas usando los controles en la parte inferior (15 equipos por página).")
    add_bullet(doc, "Menú de acciones: haga clic en el botón de acciones (⋮) de cualquier fila para ver las opciones disponibles.")

    add_body(doc, "Opciones del menú de acciones:")
    acciones = [
        ("Ver Detalle", "Abre un modal con toda la información del equipo en modo lectura."),
        ("Editar", "Abre el formulario pre-cargado para modificar los datos del equipo."),
        ("Plan de Mantenimiento", "Abre el diálogo para programar el mantenimiento: seleccione la frecuencia y la fecha de inicio."),
        ("Reporte de Entrega", "Genera el formulario de reporte de entrega del equipo a un usuario."),
        ("Ver Reportes", "Lista todos los reportes generados para el equipo."),
        ("Ver Historial", "Muestra el historial completo de cambios del equipo."),
        ("Enviar a Bodega", "Inicia el proceso de envío a almacenamiento (requiere motivo)."),
        ("Dar de Baja", "Inicia el proceso formal de baja del equipo."),
        ("Ver Hoja de Vida (HV)", "Navega a la ficha técnica completa del equipo."),
    ]
    ta = doc.add_table(rows=1, cols=2)
    ta.style = "Table Grid"
    add_table_header(ta, ["Acción", "Descripción"])
    for i, (a, d) in enumerate(acciones):
        bg = "EBF3FB" if i % 2 == 0 else None
        add_table_row(ta, [a, d], bg=bg)
    doc.add_paragraph()

    # 3.3 Crear/Editar Equipo
    add_heading(doc, "3.3 Crear o Editar un Equipo", 2)
    add_body(doc,
        "El formulario de equipo se abre como un modal. Está organizado en secciones. "
        "Complete los campos en orden:")
    pasos_form = [
        ("1. Información Básica",
         "Ingrese nombre (obligatorio), marca, modelo, serial, placa de inventario, código y año de adquisición."),
        ("2. Ubicación",
         "Especifique la ubicación general (ejemplo: Piso 2) y la ubicación específica (ejemplo: Sala de servidores)."),
        ("3. Mantenimiento",
         "Indique si el equipo requiere mantenimiento programado. Si es así, seleccione la frecuencia "
         "(mensual, trimestral, cuatrimestral, semestral o anual). El sistema calculará automáticamente las fechas."),
        ("4. Configuración de Red",
         "Indique si el equipo es gestionable (switch, router, etc.). Si lo es, ingrese el número de puertos y el direccionamiento VLAN."),
        ("5. Asignaciones",
         "Seleccione la sede (esto filtrará los servicios disponibles), el servicio, el tipo de equipo "
         "(esto determinará qué campos de hoja de vida se mostrarán) y el usuario responsable."),
        ("6. Hoja de Vida (solo al crear)",
         "Aparecen los campos técnicos habilitados para el tipo de equipo seleccionado: especificaciones "
         "de hardware, software, red, información de adquisición y observaciones. "
         "Al editar un equipo existente, esta sección no aparece; use la vista de Hoja de Vida para modificar estos datos."),
    ]
    for paso, desc in pasos_form:
        p = doc.add_paragraph()
        p.add_run(paso + ": ").bold = True
        p.add_run(desc).font.size = Pt(11)

    add_body(doc,
        "Haga clic en Guardar cuando todos los campos obligatorios estén completos. "
        "El botón estará deshabilitado si hay errores de validación. Un spinner indicará "
        "que el guardado está en proceso.")

    # 3.4 Hoja de Vida
    add_heading(doc, "3.4 Hoja de Vida del Equipo", 2)
    add_body(doc,
        "La Hoja de Vida (HV) es la ficha técnica detallada del equipo. Para acceder, "
        "seleccione 'Ver Hoja de Vida (HV)' desde el menú de acciones del equipo.")
    add_body(doc, "Secciones de la Hoja de Vida:")
    hvs = [
        ("Datos básicos", "Información general del equipo (solo lectura, se modifica desde 'Editar')."),
        ("Especificaciones técnicas", "Campos habilitados según el tipo de equipo: procesador, RAM, disco, tóner, OS, suite ofimática, IP, MAC. Haga clic en 'Editar HV' para modificar."),
        ("Información de adquisición", "Proveedor, fechas de compra e instalación, costo, contrato, tipo de adquisición."),
        ("Tipo de uso y observaciones", "Texto libre con el uso asignado al equipo y notas adicionales."),
        ("Foto del equipo", "Haga clic en 'Subir foto' para adjuntar una imagen del equipo. Puede reemplazarla o eliminarla."),
        ("Documentos adjuntos", "Haga clic en 'Agregar documento', seleccione el tipo de documento, cargue el archivo y guarde. Los documentos existentes se pueden descargar o eliminar."),
        ("Historial de ediciones", "Registro automático de todos los cambios realizados a la HV: usuario, fecha y descripción del cambio. Solo lectura."),
    ]
    for nombre, desc in hvs:
        p = doc.add_paragraph()
        p.add_run(nombre + ": ").bold = True
        p.add_run(desc).font.size = Pt(11)

    # 3.5 Enviar a Bodega
    add_heading(doc, "3.5 Enviar Equipo a Bodega", 2)
    add_body(doc,
        "Use esta función cuando un equipo sale de servicio temporalmente y necesita ser "
        "almacenado. Desde el menú de acciones del equipo seleccione 'Enviar a Bodega'.")
    add_body(doc, "Pasos del proceso:")
    add_bullet(doc, "Se abrirá un diálogo de confirmación.")
    add_bullet(doc, "Ingrese el motivo del traslado a bodega (campo obligatorio).")
    add_bullet(doc, "Haga clic en 'Confirmar'. El equipo desaparecerá de la lista activa.")
    add_bullet(doc, "Para consultar equipos en bodega, acceda desde el Card de Inventario → 'Equipos en Bodega'.")
    add_body(doc, "Reactivar desde Bodega:")
    add_bullet(doc, "Desde la vista de Equipos en Bodega, localice el equipo y seleccione 'Reactivar' en el menú de acciones.")
    add_bullet(doc, "Opcionalmente puede cambiar la ubicación del equipo al reactivarlo.")
    add_bullet(doc, "Confirme la reactivación. El equipo volverá al inventario activo.")

    # 3.6 Dar de Baja
    add_heading(doc, "3.6 Dar de Baja un Equipo", 2)
    add_body(doc,
        "La baja es un proceso formal e irreversible que retira el equipo del inventario "
        "activo. Desde el menú de acciones seleccione 'Dar de Baja'.")
    add_body(doc, "El diálogo de baja solicita la siguiente información:")
    add_bullet(doc, "Justificación detallada del motivo de la baja (campo obligatorio).")
    add_bullet(doc, "Inventario de accesorios reutilizables (cables, mouse, teclado, etc.) que quedarán en bodega.")
    add_bullet(doc, "Identificación del usuario que autoriza la baja.")
    add_bullet(doc, "Contraseña de seguridad del usuario autenticado (para confirmar la acción).")
    add_body(doc,
        "Una vez confirmada la baja, el equipo pasa al estado 'Dado de Baja'. Desde ese "
        "momento se puede descargar el Acta de Baja en PDF desde el menú de acciones "
        "('Descargar PDF de Baja').")

    # 3.7 Por Sede / Por Servicio
    add_heading(doc, "3.7 Vista por Sede y Vista por Servicio", 2)
    add_body(doc,
        "Estas vistas ofrecen perspectivas alternativas para consultar el inventario, "
        "organizando los equipos según su ubicación física (sede) o su área funcional (servicio).")
    add_body(doc, "Cómo usarlas:")
    add_bullet(doc, "Desde el Card de Inventario, haga clic en 'Por Sede' o 'Por Servicio'.")
    add_bullet(doc, "Verá una cuadrícula de tarjetas. Cada tarjeta muestra el nombre de la sede/servicio y el número de equipos asignados.")
    add_bullet(doc, "Use el filtro superior para mostrar: Todos, Solo con equipos, o Solo sin equipos.")
    add_bullet(doc, "Use la barra de búsqueda para encontrar rápidamente una sede o servicio específico.")
    add_bullet(doc, "Haga clic en una tarjeta para ver la lista de equipos de esa sede o servicio. Desde ahí tiene acceso a todas las acciones (editar, ver HV, enviar a bodega, dar de baja, etc.).")

    page_break(doc)

    # ── 4. Módulo de Parametrización SIS ──────
    add_heading(doc, "4. Módulo de Parametrización SIS", 1)
    add_body(doc,
        "El Módulo de Parametrización es accesible únicamente para el Superadministrador "
        "y el Administrador de Sistemas. Desde el dashboard haga clic en el card "
        "'Parametrización'. Este módulo tiene dos secciones principales.")

    # 4.1 Tipos de Equipo
    add_heading(doc, "4.1 Administración de Tipos de Equipo", 2)
    add_body(doc,
        "Esta sección lista todos los tipos de equipo registrados en una tabla con "
        "paginación, ordenamiento y búsqueda global.")

    add_heading(doc, "4.1.1 Crear un nuevo Tipo de Equipo", 3)
    add_bullet(doc, "Haga clic en el botón 'Nuevo Tipo de Equipo'.")
    add_bullet(doc, "Complete el formulario: nombre (obligatorio), materiales consumibles, herramientas necesarias, tiempo estimado de mantenimiento en minutos, repuestos mínimos recomendados y tipo de actividad.")
    add_bullet(doc, "En la sección 'Campos de Hoja de Vida', active o desactive los toggles de cada campo según corresponda para este tipo de equipo. Los grupos disponibles son: Red (IP, MAC), Hardware (Procesador, RAM, Disco, Tóner), Software (OS, Suite ofimática) y Gestión (Usuario, Tipo de uso, Fecha adquisición, Observaciones).")
    add_bullet(doc, "Haga clic en 'Guardar'.")

    add_heading(doc, "4.1.2 Editar un Tipo de Equipo", 3)
    add_bullet(doc, "Haga clic en el ícono de edición (lápiz) en la fila del tipo a modificar.")
    add_bullet(doc, "Modifique los campos necesarios, incluyendo los toggles de campos de hoja de vida.")
    add_bullet(doc, "Haga clic en 'Guardar'.")
    add_body(doc,
        "Importante: si cambia los campos de hoja de vida de un tipo, los cambios aplican "
        "a todos los equipos nuevos de ese tipo. Las hojas de vida ya existentes no se "
        "modifican retroactivamente.")

    add_heading(doc, "4.1.3 Activar o Desactivar un Tipo de Equipo", 3)
    add_bullet(doc, "Haga clic en el botón de estado (Activo/Inactivo) en la fila correspondiente.")
    add_bullet(doc, "Se pedirá confirmación antes de proceder.")
    add_body(doc,
        "Restricción: no es posible desactivar un tipo de equipo si tiene equipos activos "
        "asignados. Primero debe reasignar o dar de baja los equipos de ese tipo.")

    add_heading(doc, "4.1.4 Ver y Gestionar Protocolos desde la tabla", 3)
    add_bullet(doc, "Haga clic en el ícono de protocolos en la fila del tipo de equipo.")
    add_bullet(doc, "Se abrirá un modal con la lista de protocolos del tipo.")
    add_bullet(doc, "Puede agregar, editar y activar/desactivar protocolos directamente desde este modal.")

    # 4.2 Protocolos
    add_heading(doc, "4.2 Gestión de Protocolos de Mantenimiento Preventivo", 2)
    add_body(doc,
        "También puede gestionar los protocolos de mantenimiento desde la sección dedicada "
        "de Parametrización. Esta vista le permite trabajar sobre los protocolos de un tipo "
        "de equipo con más espacio y comodidad.")

    add_heading(doc, "4.2.1 Consultar protocolos de un tipo de equipo", 3)
    add_bullet(doc, "Seleccione el tipo de equipo desde el desplegable superior.")
    add_bullet(doc, "Se cargará la lista de todos los pasos del protocolo de mantenimiento preventivo para ese tipo.")
    add_bullet(doc, "La lista está paginada y muestra el texto del paso y su estado (Activo/Inactivo).")

    add_heading(doc, "4.2.2 Agregar un paso al protocolo", 3)
    add_bullet(doc, "Con el tipo de equipo seleccionado, haga clic en 'Agregar Paso'.")
    add_bullet(doc, "Ingrese el texto descriptivo del paso de mantenimiento (por ejemplo: 'Limpiar ventiladores con aire comprimido').")
    add_bullet(doc, "Haga clic en 'Guardar'. El nuevo paso aparecerá en la lista.")

    add_heading(doc, "4.2.3 Editar un paso existente", 3)
    add_bullet(doc, "Haga clic en el ícono de edición del paso que desea modificar.")
    add_bullet(doc, "Actualice el texto y haga clic en 'Guardar'.")

    add_heading(doc, "4.2.4 Activar o desactivar un paso", 3)
    add_bullet(doc, "Haga clic en el botón de estado del paso (Activo/Inactivo).")
    add_bullet(doc, "Los pasos inactivos no aparecerán al técnico durante la ejecución de un mantenimiento preventivo.")
    add_body(doc,
        "Recomendación: en lugar de eliminar pasos obsoletos, desactívelos. Así mantiene "
        "el historial de protocolos y puede reactivarlos si es necesario.")

    page_break(doc)

    # ── 5. Preguntas frecuentes ────────────────
    add_heading(doc, "5. Preguntas Frecuentes", 1)

    faqs = [
        ("¿Puedo editar los campos de la hoja de vida de un equipo ya creado?",
         "Sí. Vaya al menú de acciones del equipo, seleccione 'Ver Hoja de Vida (HV)' y haga "
         "clic en 'Editar HV'. Tenga en cuenta que los campos disponibles dependen de la "
         "configuración del tipo de equipo en Parametrización."),
        ("¿Qué diferencia hay entre Bodega y Baja?",
         "Bodega es un estado temporal: el equipo está fuera de servicio pero puede ser "
         "reactivado. Baja es definitiva: el equipo sale del inventario activo de forma "
         "permanente y se genera un acta formal."),
        ("¿Puedo cambiar el tipo de un equipo después de crearlo?",
         "Sí, desde el formulario de edición del equipo. Tenga en cuenta que cambiar el "
         "tipo puede alterar qué campos de hoja de vida son visibles."),
        ("¿Cómo exporto el inventario completo?",
         "Desde la vista de Tipos de Equipo, haga clic en el botón de exportar y seleccione "
         "'Todos'. Se descargará un archivo Excel con todos los equipos del sistema."),
        ("¿Puedo desactivar un tipo de equipo que tiene equipos asignados?",
         "No. El sistema bloqueará la desactivación si existen equipos activos de ese tipo. "
         "Primero deberá dar de baja, enviar a bodega o reasignar todos los equipos del tipo."),
        ("¿El historial de la Hoja de Vida muestra quién hizo cada cambio?",
         "Sí. El historial registra automáticamente el nombre del usuario, la fecha y hora, "
         "y el tipo de cambio realizado."),
    ]

    for pregunta, respuesta in faqs:
        p = doc.add_paragraph()
        p.add_run("P: " + pregunta).bold = True
        doc.add_paragraph("R: " + respuesta)
        doc.add_paragraph()

    # Guardar
    path = os.path.join(OUTPUT_DIR, "Manual_Usuario_Superadmin_Sistemas_HRCATCH2026.docx")
    doc.save(path)
    print(f"Manual guardado: {path}")
    return path


if __name__ == "__main__":
    p1 = generar_informe()
    p2 = generar_manual()
    print("\nDocumentos generados exitosamente:")
    print(f"  {p1}")
    print(f"  {p2}")
